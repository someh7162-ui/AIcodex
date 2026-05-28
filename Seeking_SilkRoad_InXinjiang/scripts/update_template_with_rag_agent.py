from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


WORK_DIR = Path("outputs/template_doc_update")
SRC = next(WORK_DIR.glob("*_working.docx"))
OUT = WORK_DIR / "3 软件应用与开发类作品设计和开发文档模板(1)_加入RAG和Agent.docx"


def find_para(doc, text):
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"paragraph not found: {text}")


def find_range_end(doc, start_text, next_text):
    start = next(i for i, paragraph in enumerate(doc.paragraphs) if start_text in paragraph.text)
    end = next(i for i, paragraph in enumerate(doc.paragraphs) if i > start and next_text in paragraph.text)
    return doc.paragraphs[end - 1]


def insert_after(paragraph, text="", style_source=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style_source is not None:
        new_para.style = style_source.style
        new_para.paragraph_format.space_before = style_source.paragraph_format.space_before
        new_para.paragraph_format.space_after = style_source.paragraph_format.space_after
        new_para.paragraph_format.line_spacing = style_source.paragraph_format.line_spacing
        new_para.paragraph_format.first_line_indent = style_source.paragraph_format.first_line_indent
        new_para.paragraph_format.left_indent = style_source.paragraph_format.left_indent
    if text:
        new_para.add_run(text)
    return new_para


def insert_block_after(anchor, style_para, texts):
    current = anchor
    for text in texts:
        current = insert_after(current, text, style_para)
    return current


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def set_cell_text_like(cell, text, source_cell):
    clear_cell(cell)
    source_para = source_cell.paragraphs[0] if source_cell.paragraphs else None
    paragraph = cell.paragraphs[0]
    if source_para is not None:
        paragraph.style = source_para.style
        paragraph.paragraph_format.alignment = source_para.paragraph_format.alignment
        paragraph.paragraph_format.space_before = source_para.paragraph_format.space_before
        paragraph.paragraph_format.space_after = source_para.paragraph_format.space_after
        paragraph.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    run = paragraph.add_run(text)
    if source_para is not None and source_para.runs:
        source_run = source_para.runs[0]
        run.bold = source_run.bold
        run.italic = source_run.italic
        run.underline = source_run.underline
        run.font.name = source_run.font.name
        if source_run._element.rPr is not None and source_run._element.rPr.rFonts is not None:
            east_asia = source_run._element.rPr.rFonts.get(qn("w:eastAsia"))
            if east_asia:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        run.font.size = source_run.font.size
        run.font.color.rgb = source_run.font.color.rgb


def append_row_like(table, values):
    template = table.rows[-1]
    new_tr = deepcopy(template._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for idx, value in enumerate(values):
        if idx < len(row.cells):
            set_cell_text_like(row.cells[idx], value, template.cells[idx])
    return row


def replace_paragraph_text_keep_style(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def update_table_cell(table, first_col_value, col_index, text):
    for row in table.rows:
        if row.cells[0].text.strip() == first_col_value:
            set_cell_text_like(row.cells[col_index], text, row.cells[col_index])


def main():
    doc = Document(str(SRC))

    normal_style = doc.paragraphs[63]
    h3_style = doc.paragraphs[62]
    list_style = doc.paragraphs[105]

    # 1.4 功能需求表：保留原表格样式，只增补 AI/RAG 描述。
    feature_table = doc.tables[1]
    for row in feature_table.rows:
        if row.cells[0].text.strip() == "F06":
            set_cell_text_like(row.cells[2], "支持快捷问题、上下文带入、后端 RAG 检索、结构化行程规划和来源展示", row.cells[2])
            set_cell_text_like(row.cells[3], "已升级，前端调用 /api/ai/travel-plan，由服务端统一管理 AI Key 和知识检索", row.cells[3])
    append_row_like(
        feature_table,
        [
            "F10",
            "新疆旅行 RAG + 规划 Agent",
            "从应用内目的地和攻略中检索证据，生成 3-7 天行程、交通逻辑、注意事项和可替换景点",
            "MVP 已实现，核心位于 server/ai-rag.js 与 pages/ai-assistant/index.vue",
        ],
    )

    # 3.1 技术选型表：沿用原 AI 服务行样式。
    tech_table = doc.tables[5]
    for row in tech_table.rows:
        if row.cells[0].text.strip() == "AI服务":
            set_cell_text_like(row.cells[1], "阿里云百炼兼容 OpenAI Chat Completions、后端 RAG 服务", row.cells[1])
            set_cell_text_like(row.cells[2], "模型 Key 由服务端统一管理，结合应用内知识检索降低幻觉并支撑结构化行程输出", row.cells[2])

    # 3.2.3 AI 旅游助手：在原段落后补充，不改变小节样式。
    ai_end = find_range_end(doc, "3.2.3 AI 旅游助手", "3.2.4 发现流攻略社区页")
    insert_block_after(
        ai_end,
        normal_style,
        [
            "在最新 MVP 中，AI 助手进一步升级为“新疆旅行 RAG + 规划 Agent”。前端不再直接请求大模型，而是调用后端统一的 /api/ai/travel-plan 接口，由服务端完成应用内知识检索、提示词约束、模型调用和结果解析。",
            "回答结果除了自然语言建议外，还会返回结构化 plan 与 sources。sources 用于在聊天气泡下展示“参考资料”，用户可以看到本次建议参考了哪些目的地或攻略，并可跳转回详情页继续阅读。",
        ],
    )

    # 3.4 后端接口设计表：新增接口行。
    append_row_like(
        doc.tables[7],
        ["POST", "/api/ai/travel-plan", "messages、preferences、context", "answer、plan、sources", "新疆旅行 RAG 检索与行程规划接口"],
    )

    # 3.5.3：补充服务端 RAG 机制。
    agent_end = find_range_end(doc, "3.5.3 Agent 业务上下文动态注入与边界控制", "3.5.4 图片缓存")
    current = insert_after(
        agent_end,
        "升级后的 Agent 不再只依赖前端拼接上下文，而是在服务端通过 RAG 流程动态构造证据包。系统会综合用户最近消息、景区详情页传入的 context，以及 preferences 中的天数、出发城市、区域、预算和节奏，先解析旅行意图，再检索相关知识。",
        normal_style,
    )
    for item in [
        "目的地知识来自 common/destination-data.js，包含景区名称、地区、类别、介绍、建议、季节、票价参考和开放时间参考。",
        "攻略知识来自 PostgreSQL 中已发布的 guides 与 guide_sections，优先补充真实攻略、章节正文、亮点和提示。",
        "检索结果被整理为可解释 sources，并随 AI 答案返回前端，减少“黑盒推荐”的不透明感。",
    ]:
        current = insert_after(current, item, list_style)

    # 3.5.7：新增小节，使用已有 3.5.x 标题样式与正文样式。
    bind_end = find_range_end(doc, "3.5.6 攻略与景区的结构化绑定", "3.6 主要目录说明")
    current = insert_after(bind_end, "3.5.7 新疆旅行 RAG + 规划 Agent MVP", h3_style)
    for kind, text in [
        ("normal", "本次新增的 RAG + 规划 Agent 是项目智能化能力的核心增强。它采用“轻量检索 + 受约束生成 + 结构化返回”的 MVP 架构，在不引入向量数据库的前提下，先完成可演示、可解释、可扩展的新疆旅行规划闭环。"),
        ("normal", "模型调用从前端迁移到 Node/Express 后端，服务端通过 AI_BASE_URL、AI_MODEL、AI_API_KEY 读取配置。这样可以避免在移动端保存或暴露模型 Key，也便于后续增加限流、日志、模型切换和费用控制。"),
        ("normal", "server/ai-rag.js 会把应用内目的地、攻略和通用旅行提示构造成知识块，再根据用户问题进行关键词、地区、类别、目的地别名、旅行主题和典型路线加权排序。对于“喀纳斯适合几天”“沙漠穿越装备”等问题，系统会优先命中对应目的地或户外安全知识，而不是仅依赖大模型泛化回答。"),
        ("normal", "接口要求模型返回 JSON，包括 answer、plan.days、warnings、estimatedBudget 等字段。answer 用于手机端聊天展示，plan 用于后续扩展为卡片式行程，sources 用于展示参考来源。"),
        ("bullet", "answer：面向用户的完整中文回答，包含每日安排、交通逻辑、注意事项和可替换景点。"),
        ("bullet", "plan.days：按天组织的结构化行程，每天包含主题、停靠点、交通方式和提示。"),
        ("bullet", "sources：RAG 检索命中的目的地或攻略证据，包含 sourceType、title、id、snippet 等信息。"),
        ("normal", "规划 Agent 明确不声称实时联网查票、实时天气、实时营业状态或已完成预订。缺少月份、预算、同行人群等信息时，系统使用默认值继续给出可执行方案，并提示用户补充信息后可以继续优化。AI_API_KEY 未配置时，接口返回明确错误，前端展示友好提示而不是白屏。"),
    ]:
        current = insert_after(current, text, list_style if kind == "bullet" else normal_style)

    # 3.6 目录说明表：新增文件说明。
    append_row_like(
        doc.tables[9],
        ["server/ai-rag.js", "后端 RAG 与规划 Agent 核心模块，负责偏好解析、轻量检索、提示词构造、模型调用和结构化结果解析"],
    )
    append_row_like(doc.tables[9], ["config/ai.js", "前端仅保留 AI 会话记录键和清空方法，不再保存模型 API Key"])

    # 4.2 测试用例表：新增三条测试。
    append_row_like(
        doc.tables[10],
        [
            "T15",
            "RAG 检索命中",
            "输入“第一次去新疆 5 天怎么玩”“喀纳斯适合安排几天”“沙漠穿越装备”",
            "返回相关目的地或攻略来源，如天山天池、喀纳斯、库木塔格沙漠等",
            "检索召回与证据可解释性",
        ],
    )
    append_row_like(
        doc.tables[10],
        [
            "T16",
            "行程规划接口",
            "调用 POST /api/ai/travel-plan 并传入 messages/context/preferences",
            "返回 answer、plan.days、sources 三类字段",
            "接口契约与结构化输出",
        ],
    )
    append_row_like(
        doc.tables[10],
        [
            "T17",
            "AI Key 服务端兜底",
            "后端未配置 AI_API_KEY 时调用规划接口",
            "返回清晰错误提示，不暴露密钥，不导致前端白屏",
            "安全配置与异常体验",
        ],
    )

    # 4.3 技术指标表：新增指标。
    append_row_like(doc.tables[11], ["RAG 可解释性", "AI 回答附带 sources 来源列表，用户可看到参考的目的地或攻略，并可跳转到对应详情页"])
    append_row_like(doc.tables[11], ["AI 安全边界", "规划 Agent 明确只基于应用内资料和用户上下文生成建议，不声称实时查票、实时天气或已完成预订"])

    # 4.4 已知不足：更新旧 AI Key 描述。
    for row in doc.tables[12].rows:
        if "AIKey" in row.cells[0].text or "AI Key" in row.cells[0].text:
            set_cell_text_like(row.cells[0], "AI Key 已迁移至服务端", row.cells[0])
            set_cell_text_like(row.cells[1], "前端不再保存模型 Key，但部署时必须在后端环境变量中配置 AI_API_KEY", row.cells[1])
            set_cell_text_like(row.cells[2], "后续可增加服务端限流、用量统计、模型降级和管理员配置页", row.cells[2])

    # 5.1 环境准备表 + 5.2 安装步骤：改为服务端配置。
    append_row_like(doc.tables[13], ["AI_BASE_URL、AI_MODEL、AI_API_KEY", "后端 RAG 行程规划 Agent 所需环境变量，配置在 .env 中，前端不保存模型 Key"])
    install_step = find_para(doc, "在 AI 助手页面保存 AI API Key")
    replace_paragraph_text_keep_style(install_step, "在后端 .env 中配置 AI_BASE_URL、AI_MODEL、AI_API_KEY，重启 npm run api 后即可使用 AI 助手。")
    default_note = find_para(doc, "高德和 AI 属于增强功能")
    replace_paragraph_text_keep_style(
        default_note,
        "即使未启动后端服务，用户仍可浏览首页、目的地、详情和攻略，因为前端内置了目的地和攻略数据。启动后端并配置 API 后，应用会优先使用远程数据。高德和 AI 属于增强功能，未配置 Key 时只影响地图天气路线和智能问答，不影响核心内容浏览；其中 AI Key 已迁移到后端环境变量，前端不再保存模型密钥。",
    )

    doc.core_properties.comments = "Updated on the original template style with RAG and planning Agent MVP content."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
